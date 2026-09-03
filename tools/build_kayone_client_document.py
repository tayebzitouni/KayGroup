from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Documentation"
OUT = OUT_DIR / "KAY-ONE-Dossier-Client-MVP.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(25, 38, 52)
MUTED = RGBColor(90, 104, 118)
LIGHT = "E8EEF5"
SOFT = "F4F6F9"
GREEN = RGBColor(22, 101, 52)
ORANGE = RGBColor(146, 64, 14)
RED = RGBColor(153, 27, 27)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("KAY ONE - Dossier client MVP - Page ")
    set_run_font(r, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def para(doc, text="", style=None, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def numbered(doc, text, index):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    marker = p.add_run(f"{index}. ")
    set_run_font(marker, size=10.5, color=INK)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def callout(doc, title, body, fill=SOFT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def table(doc, headers, rows, widths, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    hdr = t.rows[0].cells
    for idx, h in enumerate(headers):
        shade_cell(hdr[idx], LIGHT)
        r = hdr[idx].paragraphs[0].add_run(h)
        set_run_font(r, size=font_size, color=DARK, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            r = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(t, widths)
    doc.add_paragraph()
    return t


def cover(doc):
    para(doc, "KAY ONE", bold=True, color=DARK, size=26)
    para(doc, "Dossier client MVP", bold=True, color=BLUE, size=22)
    para(doc, "Cahier des charges organise, couverture fonctionnelle, guide d'utilisation et explication des processus financiers", color=MUTED, size=12)
    doc.add_paragraph()
    rows = [
        ("Projet", "KAY ONE - Financial & Business Operating System"),
        ("Client", "KAY Groupe"),
        ("Livraison", "MVP Windows autonome"),
        ("Version", "Client MVP 1.0.0, sans MFA"),
        ("Date", "3 septembre 2026"),
        ("Executable", "KAY-ONE-Client-MVP-1.0.0-win-x64 / KayOne.exe"),
        ("Source", "https://github.com/tayebzitouni/KayGroup.git"),
    ]
    table(doc, ["Champ", "Valeur"], rows, [1900, 7460], 10)
    callout(
        doc,
        "Message central",
        "KAY ONE n'est pas un ensemble de modules separes. Le coeur de la solution est le Transaction Engine: une seule saisie met a jour automatiquement les documents, factures, echeances, fiscalite, comptabilite, tresorerie, banque, reporting et audit.",
        "EEF6FF",
    )
    doc.add_page_break()


def executive_summary(doc):
    para(doc, "1. Resume executif", style="Heading 1")
    para(doc, "Le cahier des charges demande un systeme financier centralise capable de transformer chaque operation metier en impacts coherents, controles et tracables. La version livree est un MVP client Windows autonome: elle permet de montrer et tester le principe central sans installer une infrastructure serveur complete.")
    bullet(doc, "Fait dans le MVP: moteur transactionnel central, nouvelle operation intelligente, ventes, achats, tresorerie, banque, fiscalite, comptabilite, contrats, reporting, administration, recherche globale et audit.")
    bullet(doc, "Fait dans l'interface: les ecrans Stitch ont ete integres dans une interface professionnelle, avec navigation simple et formulaires adaptes.")
    bullet(doc, "Fait dans la securite locale: mot de passe PBKDF2, roles, perimetres societe, expiration de session, separation des taches et audit.")
    bullet(doc, "A ne pas presenter comme deja livre en production entreprise: API serveur, PostgreSQL, S3, OCR reel, connecteurs bancaires, BI externe, chiffrement centralise, PRA et supervision.")
    callout(doc, "Positionnement honnete", "Cette livraison est une version demonstrable et exploitable en local pour valider les processus metier. Elle sert de base solide pour la phase serveur/multi-utilisateurs.", "F5FBF6")


def architecture(doc):
    para(doc, "2. Vision organisee du cahier des charges", style="Heading 1")
    para(doc, "Le besoin client peut etre organise en cinq blocs au lieu de 67 points disperses.")
    rows = [
        ("Moteur central", "Une operation unique doit generer les impacts metier, fiscaux, comptables, analytiques, financiers, bancaires et audit."),
        ("Referentiels", "Societes, sites, laboratoires, clients, fournisseurs, banques, caisses, roles, comptes et regles parametrables."),
        ("Operations metier", "Achats, ventes, encaissements, decaissements, import, export, immobilisations, fiscalite, notes de frais et operations diverses."),
        ("Finance et controle", "Tresorerie, paiements, rapprochement bancaire, balance agee, journal global, reporting, rentabilite et DSO."),
        ("Industrialisation", "API, PostgreSQL, stockage S3, OCR/Document AI, IA, BI, sauvegarde, PRA, supervision et securite avancee."),
    ]
    table(doc, ["Bloc", "Ce que le client veut"], rows, [2100, 7260], 9.5)
    para(doc, "2.1 Principe de fonctionnement", style="Heading 2")
    for index, step in enumerate([
        "L'utilisateur saisit une operation dans Nouvelle operation.",
        "Le moteur lit le type, la nature, la societe, le tiers, le montant, la devise et les dimensions analytiques.",
        "Le moteur cree les objets lies: document, facture si necessaire, echeance, creance/dette, taxes, ecriture comptable, mouvement de tresorerie, operation bancaire, reporting et audit.",
        "L'utilisateur peut retrouver l'origine avec la recherche globale ou la trace de l'operation.",
    ], start=1):
        numbered(doc, step, index)


def interfaces(doc):
    para(doc, "3. Guide des interfaces livrees", style="Heading 1")
    rows = [
        ("Accueil / Tableau de bord", "Vue dirigeant: KPIs, tresorerie, creances, dettes, TVA/RAS, activite, alertes et acces rapides.", "Accueil"),
        ("Nouvelle operation", "Point d'entree principal. L'utilisateur choisit Achat, Vente, Encaissement, Decaissement, Banque, Import, Export, Immobilisation, Fiscalite, Note de frais, OD, etc.", "Operations > Nouvelle operation"),
        ("Ventes", "Suivi CA, clients, factures, proformas, exports, avoirs, acomptes et retards.", "Ventes"),
        ("Achats", "Achats locaux/etrangers, services, materiel, stock ou immobilisation, fournisseur et echeance.", "Achats"),
        ("Tresorerie", "Soldes banque/caisse, entrees/sorties, previsions et paiements.", "Tresorerie"),
        ("Banque / rapprochement", "Releves, lignes bancaires, operation bancaire, matching et annulation auditee.", "Tresorerie > Rapprochement"),
        ("Fiscalite", "Regles TVA/RAS parametrees, exonérations, historique fiscal et controles.", "Fiscalite"),
        ("Comptabilite", "Ecritures comptables equilibrees, journal global, OD et liens vers operations sources.", "Comptabilite"),
        ("Contrats", "Contrats recurrents, engagements, revisions et alertes J-90/J-30/J-7.", "Contrats"),
        ("Reporting", "Rentabilite par societe, laboratoire, site, activite, centre de cout, projet et tiers.", "Reporting"),
        ("Administration", "Referentiel Groupe, utilisateurs, roles, permissions, banques, caisses et axes analytiques.", "Administration"),
        ("Recherche globale", "Recherche par facture, client, fournisseur, reference, montant, IBAN, paiement, ecriture, document, contrat ou RAS.", "Barre de recherche / page Recherche"),
        ("Audit Trail", "Historique des creations, validations, modifications, annulations et traitements automatiques.", "Audit"),
    ]
    table(doc, ["Interface", "Role", "Ou la trouver"], rows, [1900, 5260, 2200], 8.8)


def workflows(doc):
    para(doc, "4. Comment travailler dans KAY ONE", style="Heading 1")
    para(doc, "4.1 Workflow general d'une operation", style="Heading 2")
    for index, step in enumerate([
        "Ouvrir Nouvelle operation.",
        "Choisir le type d'operation: Achat, Vente, Encaissement, Decaissement, Banque, Caisse, Import, Export, Immobilisation, Paie, Fiscalite, Note de frais, Operation diverse ou Autre.",
        "Renseigner uniquement les champs affiches par le formulaire. Si la devise est MAD, les champs de change restent caches. Si la devise est EUR/USD/GBP, le cours, la contre-valeur, les frais et les ecarts apparaissent.",
        "Ajouter la societe, le tiers, le site, le laboratoire, l'activite, le centre de cout et le projet quand ils sont connus.",
        "Importer ou declarer le justificatif. Dans le MVP, si le fichier binaire n'est pas encore stocke, le document reste en attente.",
        "Enregistrer, puis soumettre, valider, payer/rapprocher et comptabiliser selon le workflow.",
    ], start=1):
        numbered(doc, step, index)
    para(doc, "4.2 Exemple simple: achat de service local", style="Heading 2")
    table(doc, ["Action utilisateur", "Impact automatique"], [
        ("Saisir Achat > Service, fournisseur, societe, date, montant HT, TVA, centre de cout.", "Creation d'une dette fournisseur, d'une echeance, d'une ligne fiscale TVA/RAS si applicable, d'une ecriture d'achat et d'une trace d'audit."),
        ("Valider l'operation.", "Les impacts deviennent officiels dans les journaux et tableaux de bord."),
        ("Preparer le paiement.", "Creation d'un ordre de paiement lie a l'echeance."),
        ("Valider le paiement par un autre utilisateur.", "Mouvement de tresorerie et operation bancaire; la dette diminue."),
        ("Rapprocher avec le releve bancaire.", "Le paiement est confirme avec la banque et reste lie a la facture source."),
    ], [3200, 6160], 9)
    para(doc, "4.3 Exemple devise: facture fournisseur etrangere", style="Heading 2")
    para(doc, "Si une facture de 1 200 EUR est saisie, KAY ONE calcule sa contre-valeur MAD avec le cours choisi, presente le brut, la RAS eventuelle, le net fournisseur, les frais bancaires, le montant debite et l'ecart de change. L'objectif est d'eviter les calculs disperses dans Excel.")


def accounting(doc):
    doc.add_page_break()
    para(doc, "5. Comptabilite expliquee simplement", style="Heading 1")
    callout(doc, "Idee cle", "La comptabilite de KAY ONE sert a transformer une operation reelle en ecriture officielle. Une ecriture est toujours equilibree: total debit = total credit.", "FFF7ED")
    para(doc, "5.1 Debit et credit sans jargon", style="Heading 2")
    bullet(doc, "Debit: cote qui enregistre souvent ce que l'entreprise recoit ou consomme: charge, actif, argent qui entre en banque selon le cas.")
    bullet(doc, "Credit: cote qui enregistre souvent ce que l'entreprise doit ou ce qui finance l'operation: dette fournisseur, produit de vente, TVA collectee, sortie de banque selon le cas.")
    bullet(doc, "Equilibre: si une facture d'achat vaut 1 200 MAD TTC, le total des lignes au debit doit aussi faire 1 200 MAD au credit.")
    para(doc, "5.2 Trois exemples", style="Heading 2")
    table(doc, ["Cas", "Debit", "Credit", "Sens metier"], [
        ("Vente client", "Client / creance", "Chiffre d'affaires + TVA collectee", "Le client doit payer l'entreprise."),
        ("Achat fournisseur", "Charge ou immobilisation + TVA recuperable", "Fournisseur / dette", "L'entreprise doit payer le fournisseur."),
        ("Paiement fournisseur", "Fournisseur / dette", "Banque", "La dette diminue et l'argent sort de la banque."),
    ], [1800, 2600, 2600, 2360], 8.7)
    para(doc, "5.3 Pourquoi la balance agee est importante", style="Heading 2")
    para(doc, "La balance agee classe les montants non payes par anciennete: non echu, 0-30 jours, 31-60 jours, 61-90 jours, 91-120 jours, plus de 120 jours. Elle repond a une question tres concrete: qui nous doit de l'argent, qui doit etre paye, et depuis combien de temps ?")
    para(doc, "5.4 Pourquoi le rapprochement bancaire est important", style="Heading 2")
    para(doc, "Un paiement saisi dans l'application ne suffit pas. Il faut verifier que le montant apparait aussi sur le releve bancaire. Le rapprochement relie la ligne de banque au paiement ou a l'operation source. Cela evite les doublons, les oublis et les erreurs de montant.")


def coverage_summary(doc):
    para(doc, "6. Ce qui est fait et ce qui manque", style="Heading 1")
    rows = [
        ("Termine dans le MVP", "Transaction Engine, Nouvelle operation, ventes, achats, tresorerie, banque, fiscalite, comptabilite, contrats, reporting, administration, recherche, audit, balance agee clients/fournisseurs, paiements maker-checker."),
        ("Partiel mais demonstrable", "Supplier 360, Customer 360, proforma, matching bancaire exact, prevision de tresorerie, assistant financier deterministe, referentiel legal detaille."),
        ("Non livre car infrastructure", "PostgreSQL, API REST/GraphQL, stockage S3, OCR reel, connecteurs bancaires, BI/Data Warehouse, IA LLM, chiffrement centralise, supervision, PRA, signature de code."),
        ("Decision projet", "Le MVP valide le fonctionnement metier; la production entreprise necessite une phase d'industrialisation technique."),
    ]
    table(doc, ["Statut", "Contenu"], rows, [2200, 7160], 9.2)


def appendix_matrix(doc):
    doc.add_page_break()
    para(doc, "Annexe A - Mapping complet du cahier des charges", style="Heading 1")
    rows = [
        (1, "Objet integre", "Fait", "Accueil, Operations, Audit"),
        (2, "Vision Financial & Business OS", "Fait", "Dashboard, Reporting"),
        (3, "Saisie unique / impacts / tracabilite", "Fait", "Nouvelle operation, Trace"),
        (4, "Transaction Engine", "Fait", "Moteur interne"),
        (5, "Nouvelle operation centrale", "Fait", "Operations > Nouvelle operation"),
        (6, "Formulaire progressif", "Fait", "Nouvelle operation"),
        (7, "Devises et change", "Fait", "Nouvelle operation > devise"),
        (8, "Referentiel Groupe", "Partiel", "Administration > Referentiel Groupe"),
        (9, "Dimensions analytiques", "Fait", "Nouvelle operation, Reporting"),
        (10, "Identifiants internes tiers", "Fait", "Clients, Fournisseurs"),
        (11, "Supplier 360", "Partiel", "Fournisseurs 360"),
        (12, "Customer 360", "Partiel", "Clients 360"),
        (13, "Ventes", "Partiel", "Ventes"),
        (14, "Certificats d'exoneration", "Fait", "Fiscalite > Exonerations"),
        (15, "Alertes expiration", "Fait", "Dashboard, Fiscalite, Contrats"),
        (16, "Achats locaux", "Fait", "Achats, Nouvelle operation"),
        (17, "Achats de services", "Fait", "Achats"),
        (18, "Achats materiel / immobilisation", "Fait", "Immobilisations"),
        (19, "Services etrangers", "Fait", "Achats + devise + fiscalite"),
        (20, "Factures etrangeres", "Fait", "Achats, Tresorerie"),
        (21, "Dossier import", "Fait", "Imports"),
        (22, "Cout importation", "Fait", "Imports > couts"),
        (23, "Stock ou immobilisation", "Fait", "Nouvelle operation"),
        (24, "Export", "Fait", "Ventes / Export"),
        (25, "Proforma personnes physiques", "Partiel", "Ventes / Proforma"),
        (26, "Conditions de paiement", "Fait", "Tiers + echeances"),
        (27, "Contrats et engagements", "Fait", "Contrats"),
        (28, "Historique engagements", "Fait", "Contrats + Audit"),
        (29, "Notes de frais", "Fait", "Notes de frais"),
        (30, "Commissions", "Fait", "Commissions"),
        (31, "Rapprochement bancaire", "Fait", "Tresorerie > Rapprochement"),
        (32, "Automatisation rapprochement", "Partiel", "Rapprochement exact"),
        (33, "Journal global", "Fait", "Journal global"),
        (34, "Journal et audit", "Fait", "Audit"),
        (35, "Currency Engine", "Fait", "Devise / change"),
        (36, "Gains/pertes de change", "Fait", "Devise / paiement"),
        (37, "Cloture devises", "Partiel", "Donnees pretes, batch a ajouter"),
        (38, "Fiscal Engine", "Fait", "Fiscalite"),
        (39, "Regles fiscales parametrables", "Fait", "Regles fiscales"),
        (40, "Historique RAS", "Fait", "Fiscalite + Audit"),
        (41, "Tresorerie", "Fait", "Tresorerie"),
        (42, "Prevision tresorerie", "Partiel", "Dashboard, Tresorerie"),
        (43, "Preparation paiements", "Fait", "Paiements"),
        (44, "Virements internationaux", "Fait", "Paiements devise"),
        (45, "Accounting Engine", "Fait", "Comptabilite"),
        (46, "Operations diverses", "Fait", "Nouvelle operation > OD"),
        (47, "Tiers specifiques", "Partiel", "Referentiel extensible"),
        (48, "Caisses par site", "Fait", "Caisses"),
        (49, "Balance agee", "Fait", "Balance agee"),
        (50, "DSO et creances", "Fait", "Clients 360, Reporting"),
        (51, "Dashboard Direction", "Fait", "Accueil"),
        (52, "Controle de gestion", "Fait", "Reporting"),
        (53, "Modele de donnees", "Partiel", "Moteur interne"),
        (54, "Architecture technique cible", "Partiel", "Local livre; serveur a faire"),
        (55, "IA roadmap", "Phase future", "OCR/IA a brancher"),
        (56, "Assistant financier IA", "Partiel", "Assistant deterministe"),
        (57, "Securite", "Partiel", "Mot de passe, roles, session, audit"),
        (58, "Separation des taches", "Fait", "Securite & roles, Paiements"),
        (59, "Experience utilisateur", "Fait", "Interface Stitch integree"),
        (60, "Navigation", "Fait", "Menu principal"),
        (61, "Recherche globale", "Fait", "Recherche"),
        (62, "Recherche par montant", "Fait", "Recherche avancee"),
        (63, "Roadmap", "Fait", "Matrice / phases"),
        (64, "MVP recommande", "Fait", "13 cores MVP couverts localement"),
        (65, "Principes non negociables", "Fait", "Tests moteur"),
        (66, "Texte directeur", "Fait", "Vision reprise dans le moteur"),
        (67, "Architecture finale", "Partiel", "Fonctionnel local; prod serveur a faire"),
    ]
    table(doc, ["No", "Besoin client", "Etat", "Ou dans l'app / reserve"], rows, [520, 3400, 1150, 4290], 7.6)


def validation(doc):
    para(doc, "7. Recette et preuves", style="Heading 1")
    table(doc, ["Controle", "Resultat"], [
        ("Compilation Release", "0 erreur, 0 avertissement"),
        ("Tests moteur", "28/28 controles reussis"),
        ("Tests interface", "187/187 controles reussis"),
        ("Recherche source", "Aucune reference MFA/TOTP/authenticator dans source/docs/frontend"),
        ("Package client", "KAY-ONE-Client-MVP-1.0.0-win-x64.zip regenere sans MFA"),
        ("GitHub", "Dernier push sur main: Remove MFA from client MVP"),
    ], [3300, 6060], 9.2)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    cover(doc)
    executive_summary(doc)
    architecture(doc)
    interfaces(doc)
    workflows(doc)
    accounting(doc)
    coverage_summary(doc)
    validation(doc)
    appendix_matrix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
